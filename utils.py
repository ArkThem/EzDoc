from datetime import date
import docx
import re


def merge_two_dicts(x, y):
    """Given two dictionaries, merge them into a new dict as a shallow copy."""
    z = x.copy()
    z.update(y)
    return z


def get_date(variant=0):
    day_literally = {
        "1": 'первое',
        "2": 'второе',
        '3': 'третье',
        '4': 'четвертое',
        '5': 'пятое',
        '6': 'шестое',
        '7': 'седьмое',
        '8': 'восьмое',
        '9': 'девятое',
        '10': 'десятое',
        '11': 'одиннадцатое',
        '12': 'двенадцатое',
        '13': 'тринадцатое',
        '14': 'четырнадцатое',
        '15': 'пятнадцатое',
        '16': 'шестнадцатое',
        '17': 'семнадцатое',
        '18': 'восемнадцатое',
        '19': 'девятнадцатое',
        '20': 'двадцатое',
        '21': 'двадцать первое',
        '22': 'двадцать второе',
        '23': 'двадцать третье',
        '24': 'двадцать четвертое',
        '25': 'двадцать пятое',
        '26': 'двадцать шестое',
        '27': 'двадцать седьмое',
        '28': 'двадцать восьмое',
        '29': 'двадцать девятое',
        '30': 'тридцатое',
        '31': 'тридцать первое',
    }
    month_literally = {
        '01': 'января',
        '02': "февраля",
        '03': "марта",
        '04': "апреля",
        '05': "мая",
        '06': "июня",
        '07': "июля",
        '08': "августа",
        '09': "сентября",
        '10': "октября",
        '11': "ноября",
        '12': "декабря",
    }
    year_literally = {
        '2020': 'две тысячи двадцатого',
        "2021": "две тысячи двадцать первого"
    }
    if variant == 0:
        # шестнадцатое декабря две тысячи двадцатого года
        today = date.today()
        day, month, year = today.strftime("%d/%m/%Y").split("/")
        return f'{day_literally[day]} {month_literally[month]} {year_literally[year]} года'
    elif variant == 1:
        # 16.12.2020г.
        today = date.today()
        day, month, year = today.strftime("%d/%m/%Y").split("/")
        return f"{day}.{month}.{year}г."
    elif variant == 2:
        # "16" декабря 2020г.
        today = date.today()
        day, month, year = today.strftime("%d/%m/%Y").split("/")
        return f'"{day}" {month_literally[month]} {year}г.'
    pass

def get_class_input(class_name, *args, **kwargs):
    while True:
        try:
            obj = class_name(*args, **kwargs)
            return obj
        except TypeError:
            continue


def put_substring(source_str, insert_str, pos):
    return source_str[:pos]+insert_str+source_str[pos:]


def insert_before_token(docx_obj, token, insert_str):
    for par in docx_obj.paragraphs:
        if '#!stock_paragraph!#' in par.text:
            prior_paragraph = par.insert_paragraph_before('#!stock_paragraph!#')
            prior_paragraph.text = insert_str
    for par in docx_obj.paragraphs:
        if '#!stock_paragraph!#' in par.text:
            par.text = ""


def get_input(regex, message_ask="", message_error=""):
    while True:
        input_data = input(message_ask+'\n')
        if re.match(regex, input_data):
            return input_data
        else:
            print(message_error+'\n')


def doc_replace_token(doc_obj, token, replace_string, style=None):
    for p in doc_obj.paragraphs:
        if token in p.text:
            # print(token)
            p.text = p.text.replace(token, replace_string)
            if style:
                p.style = style
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_replace_token(cell, token,  replace_string)


def doc_dict_replace(doc_obj, replace_dict, style=None, bar=None):
    for p in doc_obj.paragraphs:
        for token, value in replace_dict.items():
            if token in p.text:
                # print(token)
                p.text = p.text.replace(token, replace_dict[token])
                if style:
                    p.style = style
        if bar:
            bar()
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_dict_replace(cell, replace_dict)
