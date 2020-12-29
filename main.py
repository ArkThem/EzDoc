from stock_paragraph import *
from fill_data_class import *
from functools import reduce
from alive_progress import alive_bar
from utils import *
import re
import os
import docx
from docx.shared import Pt


DEBUG = False
SAVE_PATH = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop', "Документы из программы")


client_fio = get_class_input(FIO, debug=DEBUG)
print()
client_birth_date = get_class_input(Date, _type='2', debug=DEBUG)
print()
client_birth_place = get_class_input(JustString, _type='birth_place', debug=DEBUG)
print()
clien_seria_nomer = get_class_input(SeriaNomer, debug=DEBUG)
print()
client_data_vidachi = get_class_input(Date, _type='1', debug=DEBUG)
print()
client_kem_vidan = get_class_input(JustString, _type='kem_vidan', debug=DEBUG)
print()
client_kod_kem_vidan = get_class_input(FacilityCode, debug=DEBUG)
print()
client_registration_index = get_class_input(AdrIndex, _type='1', debug=DEBUG)
print()
cliend_registration_adress = get_class_input(JustString, _type='registration_adress', debug=DEBUG)
print()
client_fact_index = get_class_input(AdrIndex, _type='2', debug=DEBUG)
print()
client_fact_adress = get_class_input(JustString, _type='fact_adress', debug=DEBUG)
print()
client_telephone = get_class_input(Telephone, debug=DEBUG)
print()

paragraph_amount = int(get_input(
    re.compile(r"^\d"),
    message_ask="Сколько вариантов акций покупаем?",
    message_error="Нужно число ввести обычное числo"
    ))

paragraphs = [StockParagraph(debug=DEBUG) for _ in range(paragraph_amount)]
summary_paragraphs = reduce(lambda a, b: a + f"\n{b}", [x.text for x in paragraphs])

dkp = docx.Document('docs/dkp.docx')
doc_dict_replace(dkp, fill_data.token_book)
insert_before_token(dkp, StockParagraph.PARAGRAPH_TOKEN, summary_paragraphs)
dkp.save(os.path.join(SAVE_PATH, f"{fill_data.token_book['#!fio_initials!#']} ДКП.docx"))
del dkp

raspiska = docx.Document('docs/raspiska.docx')
doc_dict_replace(raspiska, fill_data.token_book)
raspiska.save(os.path.join(SAVE_PATH, f"{fill_data.token_book['#!fio_initials!#']} РАСПИСКА.docx"))
del raspiska

print('Проверь совпадение пола! Слова по типу: зарегестрирован, зарегестрирована...\n')
input("Программу можно закрыть, расписка и ДКП созданы. Сделать анкету и распоряжение для регистратора?\nНажмите Enter для продолжения.")

for stock_variant in paragraphs:
    current_case_dict = merge_two_dicts(fill_data.token_book, stock_variant.token_book)
    anketa_path, rasp_path = stock_variant.AO.registrator.documents
    reg_name = stock_variant.AO.code_name
    font_settings = stock_variant.AO.registrator.font_parameters


    anketa = docx.Document(anketa_path)
    style_anketa = anketa.styles['Normal']
    font_anketa = style_anketa.font
    font_anketa.name = font_settings[0]
    font_anketa.size = font_settings[1]
    font_anketa.bold = font_settings[2]

    with alive_bar(len(anketa.paragraphs)) as bar_anketa:
        doc_dict_replace(anketa, current_case_dict, style=style_anketa, bar=bar_anketa)
    anketa.save(os.path.join(SAVE_PATH, f"{fill_data.token_book['#!fio_initials!#']} АНКЕТА {reg_name}.docx"))
    del anketa
    del style_anketa
    del font_anketa

    rasp = docx.Document(rasp_path)
    style_rasp = rasp.styles['Normal']
    font_rasp = style_rasp.font
    font_rasp.name = font_settings[0]
    font_rasp.size = font_settings[1]
    font_rasp.bold = font_settings[2]
    with alive_bar(len(rasp.paragraphs)) as bar_raspor:
        doc_dict_replace(rasp, current_case_dict, style=style_rasp, bar=bar_raspor)

    rasp.save(os.path.join(SAVE_PATH, f"{fill_data.token_book['#!fio_initials!#']} РАСПОРЯЖЕНИЕ {reg_name}.docx"))
    del rasp
    del style_rasp
    del font_rasp

if DEBUG:
    for k, v in fill_data.token_book.items():
        print(f"{k} : {v}")
    print()
    for text in paragraphs:
        print(text.text)
        print(f"{text.stock_itogo_literally}")
        print(text.stock_itogo)
